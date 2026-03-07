#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Generate Week 1 Database Applications Revision & Preparatory Work notes.

Outputs a docx with:
- Each question: your answer + improved answer + why/notes + common traps
- Quick cheat-sheets (JOIN templates, LEFT JOIN rule, HAVING vs WHERE, subquery patterns)

This file is generated from our chat patterns; adjust wording as needed.
"""

from __future__ import annotations

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "School" / "DatabaseApplications" / "Notes" / "Week1_Revision_and_Preparatory_Work_Practice_Notes_ENZH.docx"


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_code(doc: Document, code: str):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(10)


def add_bullets(doc: Document, items: list[str]):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def main():
    doc = Document()
    add_title(doc, "Database Applications — Week 1 Revision & Preparatory Work\nPractice Notes (EN + 中文)")

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} (local)")

    add_h(doc, "How to use / 怎麼用", 1)
    add_bullets(doc, [
        "EN: For each question: first read the ‘Goal’, then try writing SQL. Compare with ‘Your version’ then ‘Improved version’.",
        "中文：每題先看目標（Goal），自己先寫SQL；再比對『你的版本』與『進階版本』，最後背口訣與陷阱。",
    ])

    add_h(doc, "Quick cheat-sheet / 快速口訣", 1)
    add_bullets(doc, [
        "COUNT(*): counts rows. COUNT(col): counts non-NULL col values. / COUNT(*)算筆數；COUNT(欄位)忽略NULL。",
        "AND binds tighter than OR → use parentheses. / AND 優先於 OR → 括號保命。",
        "LIKE '%word%': contains. / LIKE '%字%': 包含。",
        "NULL check: IS NULL / IS NOT NULL (never = NULL). / 判斷NULL要用IS NULL。",
        "FULL LIST rule: the table you want fully listed must be on the LEFT of LEFT JOIN. / 要『全列誰』誰就放LEFT JOIN左邊。",
        "Aggregates in filter → HAVING (not WHERE). / 聚合條件要用HAVING不是WHERE。",
    ])

    add_h(doc, "JOIN templates / JOIN 模板", 1)
    doc.add_paragraph("Many-to-many (movie↔star via movstar):")
    add_code(doc, "SELECT s.STARNAME\nFROM movie m\nJOIN movstar ms ON ms.MVNUMB = m.MVNUMB\nJOIN star s ON s.STARNUMB = ms.STARNUMB\nWHERE m.MVTITLE = 'Manhattan';")
    doc.add_paragraph("One-to-many (director→movie):")
    add_code(doc, "SELECT d.DIRNAME, m.MVTITLE\nFROM director d\nJOIN movie m ON m.DIRNUMB = d.DIRNUMB\nWHERE d.DIRNAME = 'Kubrick, Stanley';")

    add_h(doc, "Practice Q&A / 題目練習整理", 1)

    def q(title: str, goal_en: str, goal_zh: str, your_sql: str, improved_sqls: list[str], notes: list[str]):
        add_h(doc, title, 2)
        doc.add_paragraph(f"Goal (EN): {goal_en}")
        doc.add_paragraph(f"目標（中文）：{goal_zh}")
        doc.add_paragraph("Your version / 你的版本：")
        add_code(doc, your_sql)
        doc.add_paragraph("Improved / 進階版本：")
        for s in improved_sqls:
            add_code(doc, s)
        if notes:
            doc.add_paragraph("Notes / 補充：")
            add_bullets(doc, notes)

    q(
        "Q1 — List all movies / 列出所有電影",
        "List all rows in movie table.",
        "列出movie表的所有資料。",
        "SELECT * FROM movie;",
        [
            "SELECT MVNUMB, MVTITLE, YRMDE, MVTYPE FROM movie ORDER BY MVTITLE;",
        ],
        [
            "EN: SELECT * is acceptable but less readable.",
            "中文：SELECT * 可用但不易讀；挑欄位更像報表。",
        ],
    )

    q(
        "Q2 — Titles only / 只顯示片名",
        "Display movie titles only.",
        "只顯示片名。",
        "SELECT MVTITLE FROM movie;",
        [
            "SELECT MVTITLE FROM movie ORDER BY MVTITLE;",
        ],
        [],
    )

    q(
        "Q3 — Count movies / 計算電影數量",
        "Count all movies.",
        "計算movie表總筆數。",
        "SELECT COUNT(*) FROM movie;",
        [
            "SELECT COUNT(*) AS total_movies FROM movie;",
        ],
        [
            "COUNT(*) counts rows reliably.",
        ],
    )

    q(
        "Q4 — Horror movies / 找恐怖片",
        "Filter movies by type HORROR.",
        "用MVTYPE篩選HORROR。",
        "SELECT * FROM movie WHERE MVTYPE = 'HORROR';",
        [
            "SELECT MVNUMB, MVTITLE, YRMDE, MVTYPE FROM movie WHERE MVTYPE = 'HORROR' ORDER BY MVTITLE;",
            "-- robust (if case unknown)\nSELECT MVNUMB, MVTITLE, YRMDE, MVTYPE FROM movie WHERE UPPER(MVTYPE)='HORROR';",
        ],
        [
            "If MVTYPE is already stored as HORROR, no need for UPPER().",
        ],
    )

    q(
        "Q5 — Horror with nominations / 恐怖片且提名>0",
        "Horror movies where noms > 0.",
        "恐怖片且noms>0。",
        "SELECT MVNUMB, MVTITLE, YRMDE, MVTYPE, noms FROM movie WHERE MVTYPE='HORROR' AND noms>0 ORDER BY MVTITLE;",
        [
            "SELECT MVNUMB, MVTITLE, YRMDE, MVTYPE, noms FROM movie WHERE MVTYPE='HORROR' AND noms>0 ORDER BY noms DESC, MVTITLE;",
        ],
        [
            "Trick: ORDER BY noms DESC shows most-nominated first.",
        ],
    )

    q(
        "Q6 — (Horror & noms>0) OR Comedy / 括號AND/OR",
        "Return (Horror with noms>0) OR any Comedy.",
        "回傳：恐怖且提名>0 或 任一喜劇。",
        "SELECT MVNUMB, MVTITLE, YRMDE, MVTYPE, noms FROM movie WHERE (MVTYPE='HORROR' AND noms>0) OR (MVTYPE='COMEDY') ORDER BY MVTITLE;",
        [
            "-- same logic, but parentheses prevent mistakes\nSELECT MVNUMB, MVTITLE, YRMDE, MVTYPE, noms FROM movie WHERE (MVTYPE='HORROR' AND noms>0) OR MVTYPE='COMEDY' ORDER BY MVTITLE;",
        ],
        [
            "AND has higher precedence than OR → parentheses recommended.",
        ],
    )

    q(
        "Q7 — (Horror OR Comedy) AND noms>0 / 括號位置改變",
        "Return Horror or Comedy movies, but only those with noms>0.",
        "恐怖或喜劇，但必須noms>0。",
        "SELECT MVNUMB, MVTITLE, YRMDE, MVTYPE, noms FROM movie WHERE (MVTYPE IN ('HORROR','COMEDY')) AND noms>0 ORDER BY MVTITLE;",
        [
            "-- equivalent distributive form\nSELECT MVNUMB, MVTITLE, YRMDE, MVTYPE, noms FROM movie WHERE (MVTYPE='HORROR' AND noms>0) OR (MVTYPE='COMEDY' AND noms>0) ORDER BY MVTITLE;",
        ],
        [],
    )

    q(
        "Q8 — Awards > 3 / 得獎>3",
        "Movies with AWRD > 3.",
        "找AWRD>3。",
        "SELECT MVNUMB, MVTITLE, YRMDE, MVTYPE, AWRD FROM movie WHERE AWRD > 3 ORDER BY MVTITLE;",
        [],
        [
            "Careful: question said more than three → > 3 (not > 0).",
        ],
    )

    q(
        "Q9 — Partial match (Nuremberg) / 模糊查詢 LIKE",
        "Find movies where title contains 'Nuremberg'.",
        "片名包含Nuremberg。",
        "SELECT MVNUMB, MVTITLE, YRMDE, MVTYPE, AWRD FROM movie WHERE MVTITLE LIKE '%Nuremberg%';",
        [
            "-- if case unknown per instructions\nSELECT MVNUMB, MVTITLE FROM movie WHERE (MVTITLE LIKE '%Nuremberg%' OR MVTITLE LIKE '%NUREMBERG%' OR MVTITLE LIKE '%nuremberg%');",
            "-- cleaner (SQL Server)\nSELECT MVNUMB, MVTITLE FROM movie WHERE UPPER(MVTITLE) LIKE '%NUREMBERG%';",
        ],
        [],
    )

    q(
        "Q10 — Member activity avg rentals / 平均租借次數",
        "Average NumRent over members.",
        "member表NumRent平均。",
        "SELECT AVG(NumRent) AS avg_rentals FROM member;",
        [
            "SELECT ROUND(AVG(NumRent), 2) AS avg_rentals FROM member;",
        ],
        [],
    )

    q(
        "Q11 — Join date report + AU format / 會員加入日報表",
        "List member names and join dates (AU format DD/MM/YYYY).",
        "列出會員姓名+加入日（澳洲格式）。",
        "SELECT MMBNAME, FORMAT(JOINDATE, 'dd/MM/yyyy') AS JOINDATE_AU FROM member;",
        [
            "-- alternative\nSELECT MMBNAME, CONVERT(varchar(10), JOINDATE, 103) AS JOINDATE_AU FROM member;",
        ],
        [
            "Preferred main approach: FORMAT(date, 'dd/MM/yyyy') for easy format changes.",
        ],
    )

    q(
        "Q12 — Tenure days+years + sort / 年資天數/年數",
        "Show member tenure in days and years.",
        "顯示會員年資（天/年）。",
        "SELECT MMBNAME AS [Member Name], FORMAT(JOINDATE,'dd/MM/yyyy') AS [Date Joined], DATEDIFF(day,JOINDATE,GETDATE()) AS [Number of days since joined] FROM member ORDER BY [Number of days since joined] DESC;",
        [
            "-- include years too\nSELECT MMBNAME, DATEDIFF(day,JOINDATE,GETDATE()) AS member_days, DATEDIFF(year,JOINDATE,GETDATE()) AS member_years FROM member ORDER BY member_days DESC;",
            "-- more intuitive fractional years\nSELECT MMBNAME, ROUND(DATEDIFF(day,JOINDATE,GETDATE())/365.25,2) AS member_years_approx FROM member;",
        ],
        [
            "Trap: DATEDIFF(year, ...) counts year-boundaries, not exact full years.",
        ],
    )

    q(
        "Q13 — Director+Movie list / 導演+電影清單",
        "List directors with their movies (incl. directors with no movies).",
        "列出所有導演及其電影（含沒有電影者）。",
        "SELECT d.DIRNAME, m.MVTITLE FROM director d LEFT JOIN movie m ON d.DIRNUMB=m.DIRNUMB ORDER BY d.DIRNAME, m.MVTITLE;",
        [],
        [
            "LEFT JOIN keeps all directors; missing movies appear as NULL.",
        ],
    )

    q(
        "Q14 — Actors full list with movies / 演員全清單+電影",
        "Full list of actors with their movies (NULL if none).",
        "列出所有演員及其電影（沒有則NULL）。",
        "SELECT s.STARNAME, m.MVTITLE FROM star s LEFT JOIN movstar ms ON s.STARNUMB=ms.STARNUMB LEFT JOIN movie m ON ms.MVNUMB=m.MVNUMB ORDER BY s.STARNAME, m.MVTITLE;",
        [],
        [],
    )

    q(
        "Q15 — Co-actors of Woody Allen / 同片演員（self-join / IN / EXISTS）",
        "Find co-actors who played in any movie with Woody Allen.",
        "找與Woody Allen同片演出的演員。",
        "-- preferred IN approach (template)\nSELECT DISTINCT s2.STARNAME\nFROM star s2\nJOIN movstar ms2 ON ms2.STARNUMB=s2.STARNUMB\nWHERE ms2.MVNUMB IN (\n  SELECT ms1.MVNUMB\n  FROM star s1\n  JOIN movstar ms1 ON ms1.STARNUMB=s1.STARNUMB\n  WHERE s1.STARNAME LIKE '%Allen%' AND s1.STARNAME LIKE '%Woody%'\n)\nAND NOT (s2.STARNAME LIKE '%Allen%' AND s2.STARNAME LIKE '%Woody%')\nORDER BY s2.STARNAME;",
        [
            "-- self-join\nSELECT DISTINCT s2.STARNAME\nFROM star s1\nJOIN movstar ms1 ON ms1.STARNUMB=s1.STARNUMB\nJOIN movstar ms2 ON ms2.MVNUMB=ms1.MVNUMB\nJOIN star s2 ON s2.STARNUMB=ms2.STARNUMB\nWHERE s1.STARNAME LIKE '%Allen%' AND s1.STARNAME LIKE '%Woody%'\n  AND s2.STARNAME<>s1.STARNAME\nORDER BY s2.STARNAME;",
            "-- EXISTS\nSELECT DISTINCT s2.STARNAME\nFROM star s2\nJOIN movstar ms2 ON ms2.STARNUMB=s2.STARNUMB\nWHERE EXISTS (\n  SELECT 1\n  FROM star s1\n  JOIN movstar ms1 ON ms1.STARNUMB=s1.STARNUMB\n  WHERE ms1.MVNUMB=ms2.MVNUMB\n    AND s1.STARNAME LIKE '%Allen%' AND s1.STARNAME LIKE '%Woody%'\n)\nAND NOT (s2.STARNAME LIKE '%Allen%' AND s2.STARNAME LIKE '%Woody%')\nORDER BY s2.STARNAME;",
        ],
        [
            "Rule: easiest to remember is IN (two-step).",
        ],
    )

    q(
        "Q16 — Most recent & oldest movie / 最新與最舊",
        "Find newest (MAX YRMDE) and oldest (MIN YRMDE) movies.",
        "找最新/最舊電影。",
        "SELECT MVTITLE, YRMDE FROM movie WHERE YRMDE=(SELECT MAX(YRMDE) FROM movie);\nSELECT MVTITLE, YRMDE FROM movie WHERE YRMDE=(SELECT MIN(YRMDE) FROM movie);",
        [],
        [
            "Use semicolon or run separately.",
        ],
    )

    q(
        "Q17 — Hitchcock movies not borrowed by member 4 / Hitchcock未借過",
        "List Hitchcock movies (DIRNUMB=2) not borrowed by member 4.",
        "列出Hitchcock(2)的電影中會員4未借過的。",
        "SELECT m.MVNUMB, m.MVTITLE, m.YRMDE\nFROM movie m\nWHERE m.DIRNUMB=2\n  AND NOT EXISTS (SELECT 1 FROM borrow b WHERE b.MMBNUMB=4 AND b.MVNUMB=m.MVNUMB)\nORDER BY m.MVTITLE;",
        [
            "-- NOT IN version\nSELECT m.MVNUMB, m.MVTITLE\nFROM movie m\nWHERE m.DIRNUMB=2\n  AND m.MVNUMB NOT IN (\n    SELECT b.MVNUMB\n    FROM borrow b\n    JOIN movie m2 ON m2.MVNUMB=b.MVNUMB\n    WHERE b.MMBNUMB=4 AND m2.DIRNUMB=2\n  );",
        ],
        [
            "NOT EXISTS is often safer than NOT IN when NULLs can appear.",
        ],
    )

    q(
        "Q18 — Directors report (awards) with GROUP BY/HAVING / 導演得獎報表",
        "Show directors with awards, then show all directors, then only nominated (HAVING).",
        "導演得獎/全部導演/至少提名。",
        "-- awards only\nSELECT d.DIRNAME, SUM(m.AWRD) AS total_awards\nFROM director d JOIN movie m ON d.DIRNUMB=m.DIRNUMB\nGROUP BY d.DIRNAME\nHAVING SUM(m.AWRD)>0\nORDER BY total_awards DESC;",
        [
            "-- all directors (LEFT JOIN + COALESCE)\nSELECT d.DIRNAME, COALESCE(SUM(m.AWRD),0) AS total_awards\nFROM director d LEFT JOIN movie m ON d.DIRNUMB=m.DIRNUMB\nGROUP BY d.DIRNAME\nORDER BY total_awards DESC, d.DIRNAME;",
            "-- UNION approach to replace NULL with 0\nSELECT d.DIRNAME, SUM(m.AWRD) AS total_awards\nFROM director d JOIN movie m ON d.DIRNUMB=m.DIRNUMB\nGROUP BY d.DIRNAME\nUNION\nSELECT d.DIRNAME, 0\nFROM director d LEFT JOIN movie m ON d.DIRNUMB=m.DIRNUMB\nWHERE m.DIRNUMB IS NULL\nORDER BY total_awards DESC, DIRNAME;",
            "-- nominated directors only\nSELECT d.DIRNAME, SUM(m.AWRD) AS total_awards, SUM(m.NOMS) AS total_nominations\nFROM movie m JOIN director d ON m.DIRNUMB=d.DIRNUMB\nGROUP BY d.DIRNAME\nHAVING SUM(m.NOMS) >= 1\nORDER BY total_nominations DESC, total_awards DESC;",
        ],
        [
            "WHERE filters rows; HAVING filters groups.",
        ],
    )

    doc.add_page_break()
    add_h(doc, "Space for your own notes / 你自己的筆記區", 1)
    doc.add_paragraph("…")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
